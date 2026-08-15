"""Phase 1 smoke tests for the document ingestion pipeline.

Fixtures build tiny throwaway PDF/DOCX files on the fly (via fpdf2 /
python-docx) instead of committing binary sample files to the repo — keeps
diffs clean and avoids merge conflicts on binary content.

Run: backend/.venv/Scripts/python.exe -m pytest -q
"""

from __future__ import annotations

import pytest
from docx import Document as DocxDocument
from fpdf import FPDF

from app.ingestion import (
    BlockType,
    IngestedDocument,
    SourceFormat,
    UnsupportedFormatError,
    ingest_document,
)


@pytest.fixture()
def sample_pdf_path(tmp_path):
    pdf = FPDF()

    pdf.add_page()
    pdf.set_font("Helvetica", size=14)
    pdf.cell(0, 10, "M8 Hex Bolt - Spec Sheet", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 8, "Rated Voltage: 12V. Operating Temperature: -20C to 85C.")

    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 8, "Page two body text for citation testing.")

    path = tmp_path / "sample.pdf"
    pdf.output(str(path))
    return str(path)


@pytest.fixture()
def blank_pdf_path(tmp_path):
    """A PDF with a page that has no extractable text (simulates a scan)."""
    pdf = FPDF()
    pdf.add_page()
    path = tmp_path / "blank.pdf"
    pdf.output(str(path))
    return str(path)


@pytest.fixture()
def sample_docx_path(tmp_path):
    doc = DocxDocument()
    doc.add_heading("Electrical Ratings", level=1)
    doc.add_paragraph("Rated voltage is 12V and current rating is 2A.")

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Attribute"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "IP Rating"
    table.rows[1].cells[1].text = "IP67"

    doc.add_heading("Mechanical", level=1)
    doc.add_paragraph("Overall length is 45mm.")

    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return str(path)


def test_parse_pdf_returns_ingested_document(sample_pdf_path):
    result = ingest_document(sample_pdf_path)
    assert isinstance(result, IngestedDocument)
    assert result.source_format == SourceFormat.PDF
    assert result.page_count == 2


def test_parse_pdf_preserves_page_numbers(sample_pdf_path):
    result = ingest_document(sample_pdf_path)
    text_blocks = [b for b in result.blocks if b.type == BlockType.TEXT]
    assert any(b.page == 1 for b in text_blocks)
    assert any(b.page == 2 for b in text_blocks)


def test_parse_pdf_extracts_expected_text(sample_pdf_path):
    result = ingest_document(sample_pdf_path)
    assert "Rated Voltage" in result.raw_text
    assert "Page two body text" in result.raw_text


def test_parse_pdf_flags_blank_page(blank_pdf_path):
    result = ingest_document(blank_pdf_path)
    assert result.blocks == []
    assert any("no extractable text" in w for w in result.warnings)


def test_parse_pdf_missing_file_raises():
    with pytest.raises(FileNotFoundError):
        ingest_document("/tmp/does-not-exist-12345.pdf")


def test_parse_docx_returns_ingested_document(sample_docx_path):
    result = ingest_document(sample_docx_path)
    assert isinstance(result, IngestedDocument)
    assert result.source_format == SourceFormat.DOCX
    assert result.page_count is None


def test_parse_docx_tracks_sections_from_headings(sample_docx_path):
    result = ingest_document(sample_docx_path)
    text_blocks = [b for b in result.blocks if b.type == BlockType.TEXT]
    voltage_block = next(b for b in text_blocks if "Rated voltage" in (b.text or ""))
    length_block = next(b for b in text_blocks if "Overall length" in (b.text or ""))
    assert voltage_block.section == "Electrical Ratings"
    assert length_block.section == "Mechanical"


def test_parse_docx_extracts_table_with_section(sample_docx_path):
    result = ingest_document(sample_docx_path)
    tables = result.tables
    assert len(tables) == 1
    assert tables[0].table[0] == ["Attribute", "Value"]
    assert tables[0].table[1] == ["IP Rating", "IP67"]
    assert tables[0].section == "Electrical Ratings"


def test_parse_docx_missing_file_raises():
    with pytest.raises(FileNotFoundError):
        ingest_document("/tmp/does-not-exist-12345.docx")


def test_unsupported_extension_raises():
    with pytest.raises(UnsupportedFormatError):
        ingest_document("something.txt")


def test_block_ids_are_unique_and_ordered(sample_docx_path):
    result = ingest_document(sample_docx_path)
    ids = [b.block_id for b in result.blocks]
    assert ids == sorted(ids)
    assert len(ids) == len(set(ids))
