"""DOCX parsing — Phase 1.

DOCX has no page concept in the file format itself (pagination is a rendering
concern, not stored data), so we track `section` instead: the nearest
preceding heading (any paragraph using a "Heading*" style) becomes the
section label carried on every block until the next heading. This gives
Phase 5 a citable anchor ("Installation Guide.docx, section 'Electrical
Ratings'") even without a page number.

python-docx exposes body elements as a flat sequence of paragraphs and
tables; we walk `document.element.body` directly (rather than
`document.paragraphs` / `document.tables` separately) so blocks come out in
their original document order instead of "all paragraphs, then all tables".
"""

from __future__ import annotations

import os

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.ingestion.models import BlockType, ContentBlock, IngestedDocument, SourceFormat
from app.ingestion.utils import make_block_id, normalize_whitespace, table_to_text


def _is_heading(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name or "") if paragraph.style else ""
    return style_name.lower().startswith("heading") or style_name.lower() == "title"


def parse_docx(path: str) -> IngestedDocument:
    """Parse a DOCX file into an `IngestedDocument`.

    Raises FileNotFoundError if `path` doesn't exist, and ValueError if the
    file can't be opened as a DOCX at all.
    """
    if not os.path.isfile(path):
        raise FileNotFoundError(f"No such file: {path}")

    filename = os.path.basename(path)

    try:
        document = Document(path)
    except FileNotFoundError:
        raise
    except Exception as exc:
        raise ValueError(f"Could not parse '{filename}' as a DOCX: {exc}") from exc

    blocks: list[ContentBlock] = []
    raw_text_parts: list[str] = []
    warnings: list[str] = []
    block_index = 0
    current_section: str | None = None

    body = document.element.body
    para_map = {p._element: p for p in document.paragraphs}
    table_map = {t._element: t for t in document.tables}

    for child in body.iterchildren():
        if child in para_map:
            paragraph = para_map[child]
            text = normalize_whitespace(paragraph.text or "")
            if not text:
                continue
            if _is_heading(paragraph):
                current_section = text
                blocks.append(
                    ContentBlock(
                        block_id=make_block_id(block_index),
                        type=BlockType.HEADING,
                        text=text,
                        section=current_section,
                    )
                )
            else:
                blocks.append(
                    ContentBlock(
                        block_id=make_block_id(block_index),
                        type=BlockType.TEXT,
                        text=text,
                        section=current_section,
                    )
                )
            block_index += 1
            raw_text_parts.append(text)

        elif child in table_map:
            table: Table = table_map[child]
            rows: list[list[str | None]] = []
            for row in table.rows:
                rows.append([normalize_whitespace(cell.text) or None for cell in row.cells])
            if not any(any(cell for cell in row) for row in rows):
                continue
            blocks.append(
                ContentBlock(
                    block_id=make_block_id(block_index),
                    type=BlockType.TABLE,
                    table=rows,
                    section=current_section,
                )
            )
            block_index += 1
            raw_text_parts.append(table_to_text(rows))

    if not blocks:
        warnings.append("document produced no extractable text or tables")

    return IngestedDocument(
        source_filename=filename,
        source_format=SourceFormat.DOCX,
        page_count=None,
        blocks=blocks,
        raw_text="\n\n".join(raw_text_parts),
        warnings=warnings,
    )
